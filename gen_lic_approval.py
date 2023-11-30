#!/usr/bin/env python3
# -*- coding: utf-8 -*-

"""Composes license approval for David Gast.

The script composes a DOCX file that serves as the
license approval for new and updated xbrl taxonomies.
It contains all relevant meta information and
should be submitted to the companys'
laywer David Gast every half a year.
"""

import argparse
import datetime
from docx                     import Document
from docx.enum.dml            import MSO_THEME_COLOR_INDEX
from docx.enum.text           import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.opc.constants       import RELATIONSHIP_TYPE as RT
from docx.oxml.shared         import OxmlElement, qn, CT_String
from docx.oxml.text.font      import CT_RPr
from docx.oxml.text.run       import CT_R
from docx.parts.document      import DocumentPart
from docx.shared              import Inches, Pt, RGBColor
from docx.styles.style        import _ParagraphStyle
from docx.table               import _Cell
from docx.text.run            import Run
from docx.text.paragraph      import Paragraph
from docx.text.parfmt         import ParagraphFormat
import json
from lxml.etree               import _Element
# from msilib                   import Table
import os
from typing                   import MutableMapping, Tuple
# from winreg                   import EnumValue
import xml.etree.ElementTree  as ET
from colorama                 import init
from termcolor                import colored
from Constants                import Constants

# Usage: py -3.10 gen_lic_approval.py [-family='eba'] [-version="3.2"]
#        py -3.10 gen_lic_approval.py [-family="lei"] [-version="2022-07-02 (REC)"]
#        py -3.10 gen_lic_approval.py [-family="bdp"] [-version="2.10.1 5.1"]

def add_hyperlink(paragraph: Paragraph, url: str, text: str) -> Run:
    """
    Returns an embedded hyperlink in a text string.
    The source of this code is https://github.com/python-openxml/python-docx/issues/384.
    The code has been extended with underlining blue text color.

    Keyword arguments:
    paragraph -- paragraph where text is shown
    url       -- website
    text      -- text for embedded url
    """
    part: DocumentPart = paragraph.part # get access to document.xml.rels file and new relation id value
    relation_id: str = part.relate_to(url, RT.HYPERLINK, is_external=True)
    hyperlink: _Element = OxmlElement('w:hyperlink') # create w:hyperlink tag and add new value
    hyperlink.set(qn('r:id'), relation_id)
    hyperlink.set(qn('w:history'), '1')
    new_run: CT_R = OxmlElement('w:r')

    rPr: CT_RPr = OxmlElement('w:rPr') # create w:rPr element
    rStyle: CT_String = OxmlElement('w:rStyle') # does not add hyperlink style
    rStyle.set(qn('w:val'), 'Hyperlink')
    # join all the xml elements, add required text to the w:r element
    rPr.append(rStyle)
    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)
    # create new run object and insert hyperlink
    r: Run = paragraph.add_run()
    r._r.append(hyperlink)
    # add the styling
    r.font.color.theme_color = MSO_THEME_COLOR_INDEX.HYPERLINK
    r.font.underline = True
    return r

def get_approximate_version(path_to_artifact_database: str) -> str:
    """
    Returns the latest major release version of legacy/server products as a string.
    The approximate version is the major release version of the legacy/server products.
    This version is hardcoded in the 'C:/Projects/installer/ArtifactDatabase.xml' file.

    Keyword arguments
    path_toartifact_database -- path to the 'ArtifactDatabase.xml'
    """
    try:
        tree: ET.ElementTree = ET.parse(path_to_artifact_database)
        root: ET.Element = tree.getroot()
        for elem in root.iter():
            if 'Version' in elem.tag:
                elemAttributeMap: MutableMapping[str, str] = elem.attrib.items()
                for elem_name, elem_value in elemAttributeMap:
                    if elem_name == "MajorVersionYear":
                        return elem_value
    except OSError as e:
        print("ERROR " + str(e.errno) + ": File 'C:/Projects/installer/ArtifactDatabase.xml' not found!")

def get_all_templates() -> list:
    """Return a list with all template files plus the relative path"""
    allTemplates: list = []
    for root, directories, files in os.walk("..\\..\\templates", topdown=False):
        for name in files:
            allTemplates.append(os.path.join(root, name))
    return allTemplates

"""
def iterate_over_json_file(json_file: str, elem_name: str) -> str:
    # Return requested element out of JSON file. The data are retrieved
    # from selected template of get_all_templates()
    # 
    # Keyword arguments:
    # json_file -- path to the json file
    # elem_name -- name of the element to retrieve value
    with open(json_file, "r") as data_file:
        data: dict = json.load(data_file)
        elemName: str
        elemValue: str
        for elemName, elemValue in data.items():
            if elemName == elem_name:
                return elemValue
"""

"""             
def iterate_over_license_section(json_file: str, elem_name: str) -> str:

    Return requested element out of 'license' map in template. The data
    are retrieved from selected template of get_all_templates().

    Keyword arguments:
    json_file -- path to the json file
    elem_name -- name of the element to retrieve value    

    with open(json_file, "r") as data_file:
        data: dict = json.load(data_file)
        if 'license' not in data:
            print(f"ERROR: Dict '{elem_name}' does not exist in template. Data are not available!")
        else:
            for key in data.items():
                for i,x in enumerate(key):
                    if i == 0 and x == "license":
                        for property,value in data[x].items():
                            if elem_name not in data[x]:
                                print(f"ERROR: Key '{elem_name}' does not exist in template!")
                                exit()
                            elif elem_name in data[x] and elem_name == property:
                                return value
"""
                                
def set_paragraph(header_table, row_num: int, cell_num: int, para_num: int) -> Paragraph:
    """Return a paragraph in a table cell
    
    Keyword arguments:
    header_table -- table of header secion 
    row_num      -- row in table
    cell_num     -- cell in row
    para_num     -- set paragraph in cell
    """
    table_cell: _Cell = header_table.rows[row_num].cells[cell_num]
    para_table_cell: Paragraph = table_cell.paragraphs[para_num]
    return para_table_cell

def set_pargraph_meta_section(doc_info_table, row_num: int, cell_num: int, para_num: int, format, text: str, alignment) -> _Cell:
    """
    Return a pargraph for the meta inforation section.
    This section is the very top of the license approval document.

    Keyword arguments:
    doc_info_table -- table of header secion
    row_num        -- row in table
    cell_num       -- cell in row
    para_num       -- set paragraph in cell
    """
    cell: _Cell = doc_info_table.rows[row_num].cells[cell_num]
    cell_para: Paragraph = cell.paragraphs[para_num]
    cell_para_format: ParagraphFormat = cell_para.paragraph_format
    cell_para_format.line_spacing_rule = format
    cell_para.text = text
    cell_para.alignment = alignment
    return cell

def set_title(doc: Document, format: ParagraphFormat, text: str, boldness: bool, font_size: int) -> Paragraph:
    """Return paragraph with main title contained.
    
    Keyword arguments:
    doc       -- base class
    format    -- set the format of the pargraph    
    text      -- text in pargraph
    boldness  -- set boldness of text
    font_size -- set font size for the text
    """
    title_main_obj: Paragraph = doc.add_paragraph()
    title_main_obj.paragraph_format.alignment = format
    run_main_title: Run = title_main_obj.add_run(text)
    run_main_title.bold = boldness
    run_main_title.font.size = Pt(font_size)
    return title_main_obj

def set_meta_section_table_cell_width(doc_info_section, colum_num: int, inche_num: float) -> _Cell:
    """Returns width of a table cell in meta section.

    Keyword arguments:
    doc_info_section -- section with meta informatio
    column_int       -- coumn number in table
    inche_num        -- column width
    """
    cell: _Cell
    all_cells_info_sec: Tuple[(_Cell)*3] = doc_info_section.columns[colum_num].cells 
    for cell in all_cells_info_sec:
        cell.width = Inches(inche_num)
        return cell

def set_sep_line(doc: object, line: str, boldness: bool) -> _Cell:
    """
    Returns the separation line. The line crosses the whole
    document vertically.

    Keyword arguments:
    doc      -- document object
    line     -- line in the document
    boldness -- set text bold
    """
    separation_line: _Cell = line
    sep_line_obj: _Cell = doc.add_paragraph().add_run(separation_line)
    sep_line_obj.bold = False
    return sep_line_obj

def set_main_section_paragraph(main_table, row_num: int, cell_num: int, text: str) -> Paragraph:
    """Returns one paragraph for the main section"""
    para: Paragraph = main_table.rows[row_num].cells[cell_num]
    text = ""
    para.text = text
    return para

def set_footer(footer, row_num: int, text: str, font_size: int) -> Paragraph:
    """Returns footer with text and styling"""
    footer_para: Paragraph = footer.paragraphs[row_num].add_run(text)
    footer_para.font.size = Pt(font_size)
    return footer_para

def set_additional_comment(doc: Document, alignment ,comment: str,font_size: int,rgb_color_red: int,rgb_color_yellow: int,rgb_color_green: int,italic_value: bool,bold_value: bool) -> Paragraph:
    """Returns footer with text and styling"""
    additional_comment_obj: Paragraph = doc.add_paragraph()
    additional_comment_obj.paragraph_format.alignment = alignment
    run_main_title: Run = additional_comment_obj.add_run(comment)
    run_main_title.font.size = Pt(font_size)
    run_main_title.font.color.rgb = RGBColor(rgb_color_red, rgb_color_yellow, rgb_color_green)
    run_main_title.italic = italic_value
    run_main_title.bold = bold_value
    return additional_comment_obj

def compose_docx_file_name(taxonomy_family_name: str, ws_1: str, version, ws_2: str, general_clause: str, ws_3, ph_date: str, file_extension: str ) -> str: 
    """Compose final name for the license approval daocument
    
    taxonomy_family_name: name of the taxonomy family. E.g.: "EBA
    ws_1:                 whitespace
    version:              version of the taxonomy                
    """
    composed_file_name: str = taxonomy_family_name + ws_1 + version + ws_2 + general_clause + ws_3 + ph_date + file_extension
    return composed_file_name

def main() -> None:
    """entry point"""
    argp: argparse.ArgumentParser = argparse.ArgumentParser(description='Generate license approval file to submit it to David Gast.')
    argp.add_argument('-family', '--family', help='The taxonomy\'s family name. E.g. EBA, BBK, ...')
    argp.add_argument('-version', '--version', help='The taxonomy\'s version')
    args: argparse.Namespace = argp.parse_args()

    # Initialize class with names of cell description
    objConsts: Constants = Constants()

    # Initialize modules for colors
    init()

    # set version and family name
    taxonomy_family_name: str = args.family
    taxonomy_version: str = args.version

    # Retrieve template according to family name
    if taxonomy_family_name:
        template: str = ""
        for i in range(len(get_all_templates())):
            if taxonomy_family_name.lower() in get_all_templates()[i]:
                template = get_all_templates()[i]

        # ++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++
        # <LICENSE APPROVAL DOCUMENT>
        # ++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++
        
        # Customize whole document
        doc: Document = Document()
        section_style: _ParagraphStyle = doc.styles['Normal']
        section_style.font.name = 'Calibri (Body)'
        section_style.font.size = Pt(12)
        # provide access to first section
        section = doc.sections[0]

        # ------------------------------------------------------------------------------------------------------------------
        # header section:
        # ------------------------------------------------------------------------------------------------------------------
        header: section._Header = section.header
        # table contains 1 row and 2 cells
        header_table = header.add_table(1, 2, Inches(12))
        cell: _Cell
        for cell in header_table.columns[1].cells:
            cell.width = Inches(1)
        # left cell displays 'internal usage only'
        para_l_cell: _Cell = set_paragraph(header_table, 0, 0, 0)
        run_l_cell: Run = para_l_cell.add_run(objConsts.get_header_text())
        run_l_cell.font.size = Pt(11)
        # right cell displays logo
        para_r_cell: _Cell = set_paragraph(header_table, 0, 1, 0)
        run_r_cell = para_r_cell.add_run()
        # run_r_cell.add_picture("img\\logo.png", width=1380000, height=520000)
        # set title 'THIRD PARTY SOFTWARE LICENSE APPROVAL FORM'
        set_title(doc, WD_ALIGN_PARAGRAPH.CENTER, objConsts.get_title_main_section(), True, 13)

        # ------------------------------------------------------------------------------------------------------------------
        # meta info section about the document
        # ------------------------------------------------------------------------------------------------------------------
        # create docx.document.Document object
        doc_info_section = doc.add_table(rows=3, cols=3)
        # 'From: Christoph Hartleb (Dev)'
        set_pargraph_meta_section(doc_info_section, 0, 0, 0, WD_LINE_SPACING.SINGLE, objConsts.get_sender_form(), WD_ALIGN_PARAGRAPH.LEFT)
        # 'Submitted to Legal by:'
        set_pargraph_meta_section(doc_info_section, 0, 1, 0, WD_LINE_SPACING.SINGLE, objConsts.get_submission_text_property(), WD_ALIGN_PARAGRAPH.RIGHT)
        # 'Christoph Hartleb'
        set_pargraph_meta_section(doc_info_section, 0, 2, 0, WD_LINE_SPACING.SINGLE, objConsts.get_submission_text_name(), WD_ALIGN_PARAGRAPH.LEFT)
        # 'To: David A. Gast'
        set_pargraph_meta_section(doc_info_section, 1, 0, 0, WD_LINE_SPACING.SINGLE, objConsts.get_submission_to(), WD_ALIGN_PARAGRAPH.LEFT)
        # 'Approved/Rejected by Legal:'
        set_pargraph_meta_section(doc_info_section, 1, 1, 0, WD_LINE_SPACING.SINGLE, objConsts.get_appt_or_rej_text(), WD_ALIGN_PARAGRAPH.RIGHT)
        # justify type in paragraph is left for each cell
        set_pargraph_meta_section(doc_info_section, 1, 2, 0, WD_LINE_SPACING.SINGLE, "", WD_ALIGN_PARAGRAPH.LEFT)
        # 'Submission Date:'                                                                                                               # american date format
        set_pargraph_meta_section(doc_info_section, 2, 0, 0, WD_LINE_SPACING.SINGLE, objConsts.get_sub_date() + str(datetime.datetime.now().strftime("%Y-%m-%d")), WD_ALIGN_PARAGRAPH.LEFT)
        # 'DateApproved:'
        set_pargraph_meta_section(doc_info_section, 2, 1, 0, WD_LINE_SPACING.SINGLE, objConsts.get_date_appr_text(), WD_ALIGN_PARAGRAPH.RIGHT)
        # 'YYYY-MM-DD' ->  ISO 8601 date format
        set_pargraph_meta_section(doc_info_section, 2, 2, 0, WD_LINE_SPACING.SINGLE, objConsts.get_date_format(), WD_ALIGN_PARAGRAPH.LEFT)

        set_meta_section_table_cell_width(doc_info_section, 0, 3.6)
        set_meta_section_table_cell_width(doc_info_section, 1, 3.0)
        set_meta_section_table_cell_width(doc_info_section, 2, 2.2)

        # separate header section and main section in document 
        set_sep_line(doc, "________________________________________________________________________", False)

        # ------------------------------------------------------------------------------------------------------------------
        # main section of the document (deals with meta information about the taxonomy)
        # ------------------------------------------------------------------------------------------------------------------
        main_table = doc.add_table(rows=9, cols=2)
        set_main_section_paragraph(main_table, 0, 0, objConsts.get_third_party_name_prop())

        # Name of third party software
        # ----------------------------
        col_name_h: _Cell = main_table.rows[0].cells[1]
        col_name_h.text = "xbrl taxonomy" # iterate_over_license_section(template, "swname")

        # Version number or year
        # -----------------------
        set_main_section_paragraph(main_table, 1, 0, objConsts.get_version_year_prop())
        if taxonomy_family_name == "bdp":
            # two different versions for the taxonomies provided by the Bank of Portugal.
            # therefore script call : py -3.10 gen_lic_approval.py -family="bdp" -version="2.10.1 5.0.0"
            set_main_section_paragraph(main_table, 1, 1, taxonomy_version.split(" ")[0]+" bdp v"+taxonomy_version.split(" ")[1])
        else:
            set_main_section_paragraph(main_table, 1, 1, taxonomy_version)

        # Is this a version update of 
        # previously approved software? If 
        # Yes, reason for update? 
        # --------------------------------
        set_main_section_paragraph(main_table, 2, 0, objConsts.get_update_prop())
        update_version_values: list[str] = ["Yes","No","YES","Yes, update of the ESMA ESEF Common Recommendation (CR) version"]
        if taxonomy_family_name == "dnb-dict":
            set_main_section_paragraph(main_table, 2, 1, update_version_values[1])
        elif taxonomy_family_name == "us-gaap" or taxonomy_family_name == "ifrs" or taxonomy_family_name == "xbrlgl":
            set_main_section_paragraph(main_table, 2, 1, update_version_values[2])
        elif taxonomy_family_name == "lei":
            set_main_section_paragraph(main_table, 2, 1, update_version_values[1])
        else:
            set_main_section_paragraph(main_table, 2, 1, update_version_values[0])

        # General description of software
        # -------------------------------
        set_main_section_paragraph(main_table, 3, 0, objConsts.get_softw_desc_prop())
        set_main_section_paragraph(main_table, 3, 1, "sw description")# iterate_over_license_section(template, "swdescription"))
        
        # Link to software homepage
        # -------------------------
        set_main_section_paragraph(main_table, 4, 0, objConsts.get_link_property_prop())
        homepage_hyperlink = set_paragraph(main_table, 4, 1, 0)
        if taxonomy_family_name == "us-gaap":
            add_hyperlink(homepage_hyperlink, "https://www.example.website.com", "SEC and US GAAP Taxonomies")
        elif taxonomy_family_name == "bbk":
            add_hyperlink(homepage_hyperlink, "https://www.landinpage.example.com", "Reporting - Formats(XML and XBRL)")
        elif taxonomy_family_name == "boe-banking":
            add_hyperlink(homepage_hyperlink, "https://www.landinpage.example.com", "Regulatory Reporting for the Banking Sector")
        elif taxonomy_family_name == "cipc":
            add_hyperlink(homepage_hyperlink, "https://www.landinpage.example.com", "XBRL Programs")
        elif taxonomy_family_name == "dnb-ftk":
            add_hyperlink(homepage_hyperlink, "https://www.landinpage.example.com", "Pensionsfondsen")
        elif taxonomy_family_name == "eiopa":
            add_hyperlink(homepage_hyperlink, "https://www.landinpage.example.com", "EIOPA - Tools and Data")
        elif taxonomy_family_name == "sfrdp":
            add_hyperlink(homepage_hyperlink, "https://www.example.website.com", "https://www.example.website.com")
        elif taxonomy_family_name == "acpr-corep" or taxonomy_family_name == "acpr-creditimmo" or taxonomy_family_name == "acpr-lcbft":
            add_hyperlink(homepage_hyperlink, "https://www.example.website.com", "https://www.example.website.com")
        else:
            add_hyperlink(homepage_hyperlink, "https://www.landinpage.example.com", "https://www.landinpage.example.com")

        # License type (e.g. MIT, BSD, GPL)
        # ---------------------------------
        set_main_section_paragraph(main_table, 5, 0, objConsts.get_license_prop())
        str_prep_lic_type = "license type"
        if taxonomy_family_name == "dnb-biscbs" or taxonomy_family_name == "dnb-dict" or taxonomy_family_name == "dnb-ftk":
            lic_type_hyperlink = set_paragraph(main_table, 5, 1, 0)
            add_hyperlink(lic_type_hyperlink, "license type", "CC-BY-4.0" )
        else:
            set_main_section_paragraph(main_table, 5, 1, str_prep_lic_type)

        # Link to website showing license:
        # --------------------------------
        set_main_section_paragraph(main_table, 6, 0, objConsts.get_link_lic_prop())
        licweb_hyperlink = set_paragraph(main_table, 6, 1, 0)
        if taxonomy_family_name == "us-gaap":
            add_hyperlink(licweb_hyperlink, "webpage of license", "Terms and Conditions")
        elif taxonomy_family_name == "bdp":
            add_hyperlink(licweb_hyperlink, "webpage of license", "Disclaimer and Copyright")
        elif taxonomy_family_name == "eiopa":
            add_hyperlink(licweb_hyperlink, "webpage of license", "EIOPA DPM and Taxonomy License")
        elif taxonomy_family_name == "acpr-corep" or taxonomy_family_name == "acpr-creditimmo" or taxonomy_family_name == "acpr-lcbft" or taxonomy_family_name == "bbk":
            set_main_section_paragraph(main_table, 6, 1, "webpage of license")
        else:
            add_hyperlink(licweb_hyperlink, "webpage of license", "webpage of license")
        
        if taxonomy_family_name == "boe-statistics" or taxonomy_family_name == "boe-banking" or taxonomy_family_name == "boe-insurance":
            add_hyperlink(licweb_hyperlink, iterate_over_license_section(template, "licweb1"), iterate_over_license_section(template, "licweb1"))

        # Products that will introduce license?
        # --------------------------------------------
        set_main_section_paragraph(main_table, 7, 0, objConsts.get_prod_prop())
        set_main_section_paragraph(main_table, 7, 1, objConsts.get_affected_products())

        # Approximate time/version?
        # -------------------------
        set_main_section_paragraph(main_table, 8, 0, objConsts.get_time_ver_prop())
        set_main_section_paragraph(main_table, 8, 1, (get_approximate_version('C:/Projects/installer/ArtifactDatabase.xml')))

        # ADDITIONAL COMMENTS
        # ------------------------------------------------------------------------------------------------------------------
        doc.add_paragraph().add_run("\nADDITIONAL COMMENTS:")
        if taxonomy_family_name == "dnb-biscbs" or taxonomy_family_name == "us gaap" or taxonomy_family_name == "boe-banking" or taxonomy_family_name == "cmf-cl-ci" or taxonomy_family_name == "eiopa" or taxonomy_family_name == "ifrs" or taxonomy_family_name == "acpr-corep" or taxonomy_family_name == "cipc":
            set_additional_comment(doc,WD_ALIGN_PARAGRAPH.LEFT,"Aditional comment 0",10,82,82,82,True,False)
        elif taxonomy_family_name == "eurofiling":
            set_additional_comment(doc,WD_ALIGN_PARAGRAPH.LEFT,"Additional comment",10,82,82,82,True,False)
        elif taxonomy_family_name == "us-gaap":
            set_additional_comment(doc,WD_ALIGN_PARAGRAPH.LEFT,"Aditional comment 0",8,82,82,82,True,False)
        elif taxonomy_family_name == "bbk":
            add_hyperlink(set_additional_comment(doc, WD_ALIGN_PARAGRAPH.LEFT, "", 8, 82, 82, 82, True, False),"Aditional comment 0","Aditional comment 0")
        elif taxonomy_family_name == "bdp" or taxonomy_family_name == "cbi" or taxonomy_family_name == "cbi-fsp":
            set_additional_comment(doc, WD_ALIGN_PARAGRAPH.LEFT,"Aditional comment 0",10,82,82,82,True,False)
        elif taxonomy_family_name == "edinet":
            add_hyperlink(set_additional_comment(doc, WD_ALIGN_PARAGRAPH.LEFT, "", 8, 82, 82, 82, True, False),"Aditional comment 0","Aditional comment 0")
            set_additional_comment(doc,WD_ALIGN_PARAGRAPH.LEFT,"Additional comment",11,82,82,82,True,False)
        else:
            set_additional_comment(doc,WD_ALIGN_PARAGRAPH.LEFT,"Aditional comment 0",11,82,82,82,True,False)

        # ------------------------------------------------------------------------------------------------------------------
        # footer section
        # ------------------------------------------------------------------------------------------------------------------
        footer: Paragraph = section.footer
        set_footer(footer, 0, "Ver: 01/2022", 10)

        # ++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++
        # </LICENSE APPROVAL DOCUMENT>
        # ++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++

        # Compose total filename of license approval                                    
        docx_file_name = compose_docx_file_name(
            "filebasename",
            " ",
            taxonomy_version,
            " ",
            "XBRL Taxonomy - Third Party Software License Approval Form",
            " ",
            "YYYYMMDD",
            ".docx"
        ) 

        # write content and save file
        doc.save(f"lics/{docx_file_name}")  
        print(colored("\nDocument successfully generated!", 'green')+"\n"+colored("-" * 32, 'green')+"\n"+"Your generated file: "+colored(docx_file_name, 'yellow') + " can be found at './YYYY-MM-DD/'")
    
    elif not taxonomy_family_name:
        print(f"ERROR: Taxonomy family {args.family} not found!")

if __name__ == "__main__":
    main()
